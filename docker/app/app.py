from flask import Flask, request, jsonify, send_from_directory
import os
import json
import logging
from werkzeug.utils import secure_filename
from docx import Document

# Configurar logging en un archivo
LOG_FOLDER = "/app/logs"
os.makedirs(LOG_FOLDER, exist_ok=True)
log_file = os.path.join(LOG_FOLDER, "python_app.log")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(log_file),
        logging.StreamHandler()
    ]
)

app = Flask(__name__)

UPLOAD_FOLDER = "/app/uploads"
PROCESSED_FOLDER = "/app/processed"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

@app.route('/process', methods=['POST'])
def process_document():
    try:
        if 'file' not in request.files or 'user_data' not in request.form:
            return jsonify({"error": "Faltan datos"}), 400

        file = request.files['file']
        user_data = json.loads(request.form['user_data'])

        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)

        logging.info(f"📂 Archivo recibido: {file_path}")
        logging.info(f"📌 Datos del usuario:\n{json.dumps(user_data, indent=4)}")

        processed_file_path = replace_variables(file_path, user_data)

        return jsonify({
            "message": "Documento procesado correctamente",
            "download_link": f"/download/{os.path.basename(processed_file_path)}"
        })

    except Exception as e:
        logging.error(f"❌ Error al procesar documento: {str(e)}")
        return jsonify({"error": f"Error al procesar documento: {str(e)}"}), 500


def replace_variables(file_path, user_data):
    document = Document(file_path)

    replacements = {
        "user.name": user_data.get("name", ""),
        "user.DNI": user_data.get("DNI", ""),
        "user.grade": user_data.get("grade", ""),
        "user.period": user_data.get("period", ""),
        "user.parent.name": user_data["parent"].get("name", ""),
        "user.parent.DNI": user_data["parent"].get("DNI", ""),
        "user.parent.signature": user_data["parent"].get("signature", ""),
        "institution.full_name": user_data["institution"].get("full_name", ""),
        "institution.sigla_name": user_data["institution"].get("sigla_name", ""),
        "institution.owner_name": user_data["institution"].get("owner_name", ""),
        "institution.owner_DNI": user_data["institution"].get("owner_DNI", ""),
        "institution.co_owner_name": user_data["institution"].get("co_owner_name", ""),
        "institution.co_owner_DNI": user_data["institution"].get("co_owner_DNI", ""),
        "contract.city": user_data["contract"].get("city", ""),
        "contract.day": user_data["contract"].get("day", ""),
        "contract.date": user_data["contract"].get("date", ""),
    }

    logging.info("📌 Variables a reemplazar:")
    logging.info(json.dumps(replacements, indent=4))

    detected_vars = set()

    def replace_text_safely(paragraph):
        """ Reemplaza variables sin afectar imágenes ni otros estilos. """
        full_text = "".join(run.text for run in paragraph.runs)
        modified = False
        for key, value in replacements.items():
            if f"${{{key}}}" in full_text:
                detected_vars.add(key)
                logging.info(f"🔄 Reemplazando: {key} → {value}")
                full_text = full_text.replace(f"${{{key}}}", value)
                modified = True
        if modified:
            paragraph.clear()
            paragraph.add_run(full_text)

    # 🔍 Reemplazo en texto principal
    for para in document.paragraphs:
        replace_text_safely(para)

    # 🔍 Reemplazo en encabezados y pies de página
    for section in document.sections:
        for para in section.header.paragraphs:
            replace_text_safely(para)
        for para in section.footer.paragraphs:
            replace_text_safely(para)

    # 🔍 Reemplazo en tablas
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_safely(para)

    # 🔍 Reemplazo en cuadros de texto (shapes)
    try:
        for shape in document.inline_shapes:
            if hasattr(shape, "text_frame"):
                for para in shape.text_frame.paragraphs:
                    replace_text_safely(para)
    except Exception as e:
        logging.warning(f"⚠️ No se pudieron procesar cuadros de texto: {str(e)}")

    logging.info(f"🔍 Variables detectadas en el documento: {detected_vars}")

    processed_file_path = os.path.join(PROCESSED_FOLDER, "processed_" + os.path.basename(file_path))
    document.save(processed_file_path)

    logging.info(f"✅ Documento guardado en: {processed_file_path}")
    return processed_file_path


@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(PROCESSED_FOLDER, filename, as_attachment=True)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
